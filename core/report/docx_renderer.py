"""DOCX report renderer -- for engagement writeups/deliverables that need a Word document.

Mirrors the same content structure as html_renderer.py (executive
summary, session details, category breakdown, per-category findings
tables) using python-docx. The import is local to render_docx() rather
than top-level so that JSON/HTML report generation never depends on
python-docx being installed -- only a caller that actually requests a
.docx report pays for that import (and gets a clear ImportError with
install instructions if it's missing).
"""

from __future__ import annotations

from io import BytesIO

from core.report.model import CategorySummary, InterpretedFinding, SessionReport

_SEVERITY_COLOR = {
    "high": "D03B3B",
    "medium": "FAB219",
    "low": "898781",
    "info": "898781",
}
_SEVERITY_LABEL = {"high": "HIGH", "medium": "MEDIUM", "low": "LOW", "info": "INFO"}


def _fmt_duration(seconds: float) -> str:
    seconds = int(seconds)
    if seconds < 60:
        return f"{seconds}s"
    minutes, secs = divmod(seconds, 60)
    if minutes < 60:
        return f"{minutes}m {secs}s"
    hours, minutes = divmod(minutes, 60)
    return f"{hours}h {minutes}m"


def render_docx(report: SessionReport) -> bytes:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as exc:  # pragma: no cover -- environment-dependent
        raise ImportError(
            "DOCX report generation requires python-docx. Install it with: pip install python-docx"
        ) from exc

    meta = report.meta
    doc = Document()

    # -- Title --------------------------------------------------------
    title = doc.add_heading("Frida-ShieldBreaker Session Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    subtitle.add_run(f"Target: {meta.target}    Device: {meta.device_name}    Generated: {meta.started_at.strftime('%Y-%m-%d %H:%M UTC')}").italic = True

    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run("For use only on targets you own or are explicitly authorized to test.")
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(0x89, 0x87, 0x81)

    # -- Executive summary ---------------------------------------------
    doc.add_heading("Executive Summary", level=1)
    severity_counts = report.severity_counts()
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = "Light Grid Accent 1"
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = summary_table.rows[0].cells
    for cell, label in zip(hdr, ["Total findings", "Bypassed", "Modules run", "Duration"]):
        cell.text = label
    values_row = summary_table.add_row().cells
    for cell, value in zip(
        values_row,
        [
            str(report.total_findings),
            str(report.total_bypassed),
            str(len(meta.modules_enabled)),
            _fmt_duration(meta.duration_seconds),
        ],
    ):
        cell.text = value

    doc.add_paragraph(
        f"{severity_counts.get('high', 0)} high-severity, {severity_counts.get('medium', 0)} medium, "
        f"{severity_counts.get('low', 0)} low, and {severity_counts.get('info', 0)} informational finding(s) "
        f"were recorded across {len(meta.modules_enabled)} module(s)."
    )

    # -- Session details -------------------------------------------------
    doc.add_heading("Session Details", level=1)
    details = [
        ("Target", meta.target),
        ("Device", meta.device_name),
        ("Platform / architecture", f"{meta.platform or 'unknown'} / {meta.arch or 'unknown'}"),
        ("Started", meta.started_at.isoformat()),
        ("Duration", _fmt_duration(meta.duration_seconds)),
        ("Modules requested", ", ".join(meta.modules_requested) or ("auto-detected" if meta.auto_detect else "none")),
        ("Modules enabled", ", ".join(meta.modules_enabled) or "none"),
        ("Automatic framework detection", "enabled" if meta.auto_detect else "disabled"),
        ("Active bypass requested", "yes" if meta.bypass_requested else "no"),
        ("Frida-ShieldBreaker version", meta.tool_version),
    ]
    details_table = doc.add_table(rows=0, cols=2)
    details_table.style = "Light List Accent 1"
    for key, value in details:
        row = details_table.add_row().cells
        row[0].text = key
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(value)

    # -- Findings by category ------------------------------------------
    doc.add_heading("Findings by Category", level=1)
    categories: list[CategorySummary] = sorted(report.by_category(), key=lambda c: c.total, reverse=True)
    if not categories:
        doc.add_paragraph("No findings were recorded during this session.")
    else:
        cat_table = doc.add_table(rows=1, cols=3)
        cat_table.style = "Light Grid Accent 1"
        for cell, label in zip(cat_table.rows[0].cells, ["Category", "Findings", "Bypassed"]):
            cell.text = label
        for cat in categories:
            row = cat_table.add_row().cells
            row[0].text = cat.title
            row[1].text = str(cat.total)
            row[2].text = str(cat.bypassed)

    # -- Detailed findings, one section + table per category ----------
    doc.add_heading("Detailed Findings", level=1)
    for cat in categories:
        findings: list[InterpretedFinding] = report.findings_in_category(cat.category)
        doc.add_heading(f"{cat.title} ({cat.total})", level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light List Accent 1"
        for cell, label in zip(table.rows[0].cells, ["Severity", "Finding", "Module", "Time"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True

        for f in findings:
            row = table.add_row().cells

            sev_run = row[0].paragraphs[0].add_run(_SEVERITY_LABEL.get(f.severity, f.severity.upper()))
            sev_run.bold = True
            sev_run.font.color.rgb = RGBColor.from_string(_SEVERITY_COLOR.get(f.severity, "898781"))

            finding_para = row[1].paragraphs[0]
            title_run = finding_para.add_run(f.title)
            title_run.bold = True
            row[1].add_paragraph(f.description)

            row[2].text = f.finding.module
            row[3].text = f.finding.timestamp.strftime("%H:%M:%S.%f")[:-3]

    # -- Footer ----------------------------------------------------------
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"Generated by Frida-ShieldBreaker {meta.tool_version}.")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x89, 0x87, 0x81)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
